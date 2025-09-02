from docx import Document

def generate_contract(template_path, context, clauses, styles, output_path):
    doc = Document(template_path)

    # Substituir placeholders simples
    for p in doc.paragraphs:
        for key, val in context.items():
            if f"{{{{ {key} }}}}" in p.text:
                p.text = p.text.replace(f"{{{{ {key} }}}}", str(val))

    # Inserir cláusulas no marcador {{ clausulas }}
    for p in doc.paragraphs:
        if "{{ clausulas }}" in p.text:
            p.text = p.text.replace("{{ clausulas }}", "")
            for idx, c in enumerate(clauses, start=1):
                para = doc.add_paragraph(f"Cláusula {idx}ª — {c['titulo']}
{c['texto']}")
                apply_style(para, styles["clausulas_extra"])

    doc.save(output_path)

def apply_style(paragraph, style_conf):
    run = paragraph.runs[0]
    if style_conf.get("bold"): run.bold = True
    if style_conf.get("italic"): run.italic = True
    if style_conf.get("caps"): run.text = run.text.upper()
    run.font.name = style_conf.get("font", "Times New Roman")
